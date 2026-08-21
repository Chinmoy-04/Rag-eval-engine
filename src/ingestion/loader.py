"""Load and chunk documents from data/raw_docs/ via LlamaIndex."""

from __future__ import annotations

import hashlib
import logging
from pathlib import Path

from llama_index.core import Document, SimpleDirectoryReader
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core.schema import TextNode

from src.config import AppConfig, RAW_DOCS_DIR

logger = logging.getLogger("rag_eval")

# Text / PDF via SimpleDirectoryReader; CSV / DOCX via custom readers below.
SIMPLE_READER_EXTS = {".txt", ".md", ".pdf"}
CUSTOM_READER_EXTS = {".csv", ".docx"}
SUPPORTED_EXTS = SIMPLE_READER_EXTS | CUSTOM_READER_EXTS


def list_corpus_files(docs_dir: Path | None = None) -> list[Path]:
    """Return supported corpus files, sorted for stable fingerprints."""
    root = docs_dir or RAW_DOCS_DIR
    if not root.exists():
        return []
    files = [
        path
        for path in root.rglob("*")
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTS
    ]
    return sorted(files)


def corpus_fingerprint(docs_dir: Path | None = None) -> str:
    """Hash file names, sizes, and mtimes so we can skip unchanged ingests."""
    hasher = hashlib.sha256()
    for path in list_corpus_files(docs_dir):
        stat = path.stat()
        hasher.update(path.name.encode("utf-8"))
        hasher.update(str(stat.st_size).encode("ascii"))
        hasher.update(str(int(stat.st_mtime)).encode("ascii"))
    return hasher.hexdigest()


def _base_metadata(path: Path, file_type: str) -> dict:
    return {
        "file_name": path.name,
        "file_path": str(path.resolve()),
        "file_type": file_type,
    }


def _load_csv_document(path: Path) -> Document:
    """Turn a CSV into readable text (header + rows) for embedding/retrieval."""
    raw = path.read_text(encoding="utf-8-sig").strip()
    if not raw:
        return Document(text="", metadata=_base_metadata(path, "csv"))

    lines = [line for line in raw.splitlines() if line.strip()]
    header = lines[0] if lines else ""
    body_lines = lines[1:] if len(lines) > 1 else []

    parts = [
        f"CSV table: {path.name}",
        f"Columns: {header}",
        "",
    ]
    for i, row in enumerate(body_lines, start=1):
        parts.append(f"Row {i}: {row}")

    text = "\n".join(parts).strip() + "\n"
    return Document(text=text, metadata=_base_metadata(path, "csv"))


def _load_docx_document(path: Path) -> Document:
    """Extract paragraph (+ table cell) text from a Word document."""
    try:
        from docx import Document as DocxFile
    except ImportError as exc:
        raise ImportError(
            "python-docx is required to ingest .docx files. Run: uv add python-docx"
        ) from exc

    doc = DocxFile(str(path))
    chunks: list[str] = []
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            chunks.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [((cell.text or "").strip()) for cell in row.cells]
            if any(cells):
                chunks.append(" | ".join(cells))

    body = "\n\n".join(chunks).strip()
    text = f"Word document: {path.name}\n\n{body}\n" if body else ""
    return Document(text=text, metadata=_base_metadata(path, "docx"))


def _load_custom_documents(files: list[Path]) -> list[Document]:
    documents: list[Document] = []
    for path in files:
        suffix = path.suffix.lower()
        if suffix == ".csv":
            documents.append(_load_csv_document(path))
        elif suffix == ".docx":
            documents.append(_load_docx_document(path))
    return documents


def load_documents(docs_dir: Path | None = None) -> list[Document]:
    """Load supported corpus files (.txt, .md, .pdf, .csv, .docx)."""
    root = docs_dir or RAW_DOCS_DIR
    files = list_corpus_files(root)
    if not files:
        raise FileNotFoundError(
            f"No supported documents in {root} "
            f"({', '.join(sorted(SUPPORTED_EXTS))}). "
            "Run: uv run python scripts/seed_sample_corpus.py"
        )

    logger.info("Loading %d documents from %s", len(files), root)

    simple_files = [p for p in files if p.suffix.lower() in SIMPLE_READER_EXTS]
    custom_files = [p for p in files if p.suffix.lower() in CUSTOM_READER_EXTS]

    documents: list[Document] = []
    if simple_files:
        reader = SimpleDirectoryReader(
            input_dir=str(root),
            required_exts=sorted(SIMPLE_READER_EXTS),
            recursive=True,
            filename_as_id=True,
            exclude_hidden=True,
        )
        documents.extend(reader.load_data())

    documents.extend(_load_custom_documents(custom_files))

    # Drop empty extracts so they don't pollute the index.
    documents = [doc for doc in documents if (doc.text or "").strip()]
    logger.info("Loaded %d LlamaIndex documents", len(documents))
    return documents


def chunk_documents(
    documents: list[Document],
    chunk_size: int,
    chunk_overlap: int,
) -> list[TextNode]:
    """Split documents into overlapping sentence-aware chunks.

    chunk_size / chunk_overlap are in *tokens* (LlamaIndex SentenceSplitter),
    not characters. Overlap keeps a fact that sits on a chunk boundary
    retrievable from either neighboring chunk.
    """
    splitter = SentenceSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    nodes = splitter.get_nodes_from_documents(documents)
    logger.info(
        "Chunked %d documents into %d nodes (size=%d, overlap=%d)",
        len(documents),
        len(nodes),
        chunk_size,
        chunk_overlap,
    )
    return nodes


def load_and_chunk(config: AppConfig, docs_dir: Path | None = None) -> list[TextNode]:
    """Load the corpus and return chunk nodes using config chunk settings."""
    documents = load_documents(docs_dir)
    return chunk_documents(documents, config.chunk_size, config.chunk_overlap)
